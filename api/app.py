import os
import uuid
import torch
import re
import base64
from flask import Flask, request, jsonify, send_from_directory
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from docx import Document
from docx.oxml import OxmlElement
import pdfplumber
import io

HERE = os.path.dirname(__file__)

# If the new React frontend is built, serve it as an SPA.
# This repo currently builds to `dist/apps/web` (see `web/package.json`), but we also
# support the conventional `web/dist` to keep local workflows simple.
_WEB_DIST_CANDIDATES = [
    os.path.abspath(os.path.join(HERE, "..", "dist", "apps", "web")),
    os.path.abspath(os.path.join(HERE, "..", "web", "dist")),
]
WEB_DIST_DIR = next((p for p in _WEB_DIST_CANDIDATES if os.path.isdir(p)), _WEB_DIST_CANDIDATES[0])
WEB_ASSETS_DIR = os.path.join(WEB_DIST_DIR, "assets")

app = Flask(
    __name__,
    static_folder=WEB_DIST_DIR if os.path.isdir(WEB_DIST_DIR) else None,
    static_url_path="",
)
DOWNLOAD_DIR = os.path.join(os.path.dirname(__file__), "downloads")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

# -------------------------------------------------------
# ПУТИ К MERGED МОДЕЛЯМ
# -------------------------------------------------------
MODEL_KK_RU = "/home/user/nllb_finetune_kaz_to_rus/nllb-kazakh-to-russian-merged"
MODEL_RU_KK = "/home/user/nllb_finetune_rus_to_kaz/nllb-russian-to-kazakh-merged"

BASE_MODEL = "facebook/nllb-200-distilled-1.3B"

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
print(f"[INFO] Устройство: {device}")

print("[INFO] Загружаю токенизатор...")
tokenizer = AutoTokenizer.from_pretrained(BASE_MODEL)

print("[INFO] Загружаю модель КАЗ→РУС...")
model_kk_ru = AutoModelForSeq2SeqLM.from_pretrained(MODEL_KK_RU)
model_kk_ru = model_kk_ru.to(device)
model_kk_ru.eval()

print("[INFO] Загружаю модель РУС→КАЗ...")
model_ru_kk = AutoModelForSeq2SeqLM.from_pretrained(MODEL_RU_KK)
model_ru_kk = model_ru_kk.to(device)
model_ru_kk.eval()

print("[INFO] Всё готово! Сервер запускается...")

def translate_chunk(text, model, src_lang, tgt_lang):
    # Разбиваем текст на части: КАПС фразы и обычный текст
    parts = re.split(r'(\b(?:[А-ЯЁҚҒҮҰҢӘІӨA-Z]{2,}|[А-ЯЁҚҒҮҰҢӘІӨA-Z]{1}\b)(?:\s+(?:[А-ЯЁҚҒҮҰҢӘІӨA-Z]{2,}|[А-ЯЁҚҒҮҰҢӘІӨA-Z]{1}\b))*\s*)', text)

    result = []
    for part in parts:
        if not part.strip():
            result.append(part)
            continue
        
        is_caps = bool(re.match(r'^(?:[А-ЯЁҚҒҮҰҢӘІӨA-Z]{2,}|[А-ЯЁҚҒҮҰҢӘІӨA-Z]{1}\b)(?:\s+(?:[А-ЯЁҚҒҮҰҢӘІӨA-Z]{2,}|[А-ЯЁҚҒҮҰҢӘІӨA-Z]{1}\b))*\s*$', part.strip()))

        # Переводим каждую часть
        tokenizer.src_lang = src_lang
        text_to_translate = part.strip().capitalize() if is_caps else part.strip()
        inputs = tokenizer(text_to_translate, return_tensors="pt", padding=True, truncation=True, max_length=256).to(device)
        tgt_lang_id = tokenizer.convert_tokens_to_ids(tgt_lang)
        with torch.no_grad():
            output = model.generate(
                **inputs,
                forced_bos_token_id=tgt_lang_id,
                max_length=256,
                repetition_penalty=1.3,
                no_repeat_ngram_size=4,
                num_beams=4,
            )
        translated_part = tokenizer.decode(output[0], skip_special_tokens=True)

        # Если оригинал был КАПС — возвращаем КАПС
        if is_caps:
            translated_part = translated_part.upper()

        result.append(translated_part)

    return " ".join(p for p in result if p.strip())

def translate_text(text, tgt_lang_code):
    """Разбивает текст на куски и переводит по частям."""
    if tgt_lang_code == "ru":
        model = model_kk_ru
        src_lang = "kaz_Cyrl"
        tgt_lang = "rus_Cyrl"
    else:
        model = model_ru_kk
        src_lang = "rus_Cyrl"
        tgt_lang = "kaz_Cyrl"

    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    chunks = []
    current = ""
    for sentence in sentences:
        if len(current) + len(sentence) < 200:
            current += (" " if current else "") + sentence
        else:
            if current:
                chunks.append(current)
            current = sentence
    if current:
        chunks.append(current)

    translated_parts = []
    for chunk in chunks:
        if chunk.strip():
            translated_parts.append(translate_chunk(chunk, model, src_lang, tgt_lang))

    return " ".join(translated_parts)


def translate_docx(file_bytes, tgt_lang):
    """Переводит DOCX изменяя оригинал напрямую — форматирование сохраняется."""
    doc = Document(io.BytesIO(file_bytes))

    def translate_paragraph(para):
        if not para.text.strip():
            return

        # Собираем сегменты между мягкими переносами строк <w:br/>
        segments = []
        current_segment = []
        for run in para.runs:
            for child in run._r:
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag == 't':
                    current_segment.append(('text', child.text or '', run))
                elif tag == 'br':
                    segments.append(current_segment)
                    current_segment = []
        if current_segment:
            segments.append(current_segment)

        # Переводим каждый сегмент отдельно
        translated_segments = []
        for segment in segments:
            segment_text = "".join(t for _, t, _ in segment).strip()
            if segment_text:
                translated_segments.append(translate_text(segment_text, tgt_lang))
            else:
                translated_segments.append("")

        # Очищаем все runs от текста и переносов
        for run in para.runs:
            for child in list(run._r):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag in ('t', 'br'):
                    run._r.remove(child)

        # Записываем переведённый текст в первый run
        if para.runs:
            first_r = para.runs[0]._r
            for idx, translated in enumerate(translated_segments):
                t_elem = OxmlElement('w:t')
                t_elem.text = translated
                if translated.startswith(' ') or translated.endswith(' '):
                    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                first_r.append(t_elem)
                if idx < len(translated_segments) - 1:
                    br_elem = OxmlElement('w:br')
                    first_r.append(br_elem)

    for para in doc.paragraphs:
        translate_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    translate_paragraph(para)

    filename = f"translated_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(DOWNLOAD_DIR, filename)
    doc.save(filepath)
    return filename


def extract_txt(file_bytes):
    return file_bytes.decode("utf-8", errors="ignore")


def extract_pdf(file_bytes):
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text += (page.extract_text() or "") + "\n"
    return text.strip()


def save_docx(text):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    filename = f"translated_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(DOWNLOAD_DIR, filename)
    doc.save(filepath)
    return filename


def _add_cors_headers(resp):
    resp.headers["Access-Control-Allow-Origin"] = request.headers.get("Origin", "*")
    resp.headers["Vary"] = "Origin"
    resp.headers["Access-Control-Allow-Credentials"] = "true"
    resp.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization"
    resp.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    return resp


@app.after_request
def after_request(resp):
    return _add_cors_headers(resp)


@app.route("/hcgi/api/health", methods=["GET"])
def health():
    return jsonify({"ok": True})


def _resolve_direction(source_language: str, target_language: str) -> str:
    src = (source_language or "").strip().lower()
    tgt = (target_language or "").strip().lower()

    # Supported: kk <-> ru only (models already trained)
    if src == "kk" and tgt == "ru":
        return "ru"  # translate_text expects target lang code
    if src == "ru" and tgt == "kk":
        return "kk"

    # Backwards-compat: if only target provided
    if tgt in ("kk", "ru"):
        return tgt

    raise ValueError("Unsupported language pair. Supported: kk↔ru.")


@app.route("/hcgi/api/translate", methods=["POST", "OPTIONS"])
def translate_api():
    if request.method == "OPTIONS":
        return _add_cors_headers(jsonify({"ok": True}))

    source_language = request.form.get("sourceLanguage", "")
    target_language = request.form.get("targetLanguage", "")
    try:
        tgt_lang = _resolve_direction(source_language, target_language)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    # --- File translation (multipart: file) ---
    if "file" in request.files and request.files["file"] and request.files["file"].filename:
        uploaded = request.files["file"]
        original_name = uploaded.filename
        lowered = original_name.lower()
        file_bytes = uploaded.read()

        try:
            if lowered.endswith(".docx"):
                saved_name = translate_docx(file_bytes, tgt_lang)
                saved_path = os.path.join(DOWNLOAD_DIR, saved_name)
                with open(saved_path, "rb") as f:
                    out_bytes = f.read()
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                out_name = os.path.splitext(original_name)[0] + "_translated.docx"
            elif lowered.endswith(".txt"):
                translated_text = translate_text(extract_txt(file_bytes), tgt_lang)
                out_bytes = translated_text.encode("utf-8")
                mime_type = "text/plain; charset=utf-8"
                out_name = os.path.splitext(original_name)[0] + "_translated.txt"
            elif lowered.endswith(".pdf"):
                translated_text = translate_text(extract_pdf(file_bytes), tgt_lang)
                out_bytes = translated_text.encode("utf-8")
                mime_type = "text/plain; charset=utf-8"
                out_name = os.path.splitext(original_name)[0] + "_translated.txt"
            else:
                return jsonify({"error": "Unsupported file format. Supported: PDF, DOCX, TXT."}), 400

            return jsonify({
                "translatedFile": base64.b64encode(out_bytes).decode("ascii"),
                "fileName": out_name,
                "mimeType": mime_type,
            })
        except Exception as e:
            return jsonify({"error": f"Translation error: {str(e)}"}), 500

    # --- Text translation (multipart: inputText) ---
    input_text = request.form.get("inputText", "").strip()
    if not input_text:
        return jsonify({"error": "Either inputText or file must be provided"}), 400

    try:
        translated = translate_text(input_text, tgt_lang)
        return jsonify({"translatedText": translated})
    except Exception as e:
        return jsonify({"error": f"Translation error: {str(e)}"}), 500


@app.route("/export-docx", methods=["POST"])
def export_docx():
    text = request.form.get("translated_text", "").strip()
    if not text:
        return jsonify({"message": "Нет текста для экспорта."}), 400
    try:
        saved = save_docx(text)
        return jsonify({"download_url": f"/downloads/{saved}"})
    except Exception as e:
        return jsonify({"message": str(e)}), 500


@app.route("/downloads/<filename>")
def download_file(filename):
    return send_from_directory(DOWNLOAD_DIR, filename, as_attachment=True)


def _serve_spa():
    index_path = os.path.join(WEB_DIST_DIR, "index.html")
    if not os.path.isfile(index_path):
        return (
            "Frontend is not built yet. Run the React app from `web/` (dev) or build it (`npm run build`) to serve it via Flask.",
            500,
        )
    return send_from_directory(WEB_DIST_DIR, "index.html")


@app.route("/", defaults={"path": ""})
@app.route("/<path:path>")
def spa_routes(path):
    # Serve built assets directly if Flask static is enabled
    if app.static_folder and path.startswith("assets/"):
        asset_rel = path[len("assets/"):]
        return send_from_directory(WEB_ASSETS_DIR, asset_rel)
    # React router handles /, /translate, /about, /contact, etc.
    return _serve_spa()


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=False)
